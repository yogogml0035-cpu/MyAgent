"""Task-scoped resource execution tools.

This module keeps the LangChain tool surface thin.  The tools expose stable
resource-oriented names and delegate the actual work to a local execution
adapter that mirrors a future Provision/Execute boundary.
"""

from __future__ import annotations

import json
import math
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from langchain_core.tools import BaseTool, tool

from app.storage import (
    UPLOAD_FORMATS,
    build_upload_resource_ref,
    file_sha256,
    normalize_artifact_name,
    source_format_for_upload,
)

RESOURCE_TOOL_SYSTEM_PROMPT = """上传文件属于任务资源，不是聊天上下文。
当用户询问上传文件时，优先使用资源工具。先检查或列出资源，再只读取所需的文本页或表格范围。不要假设上传文件内容已经直接出现在当前对话中。
读取 Word/docx 内容时，使用 inspect_resource、read_resource_text 或 read_resource_table。
生成 Word/docx 交付文件时，必须使用 create_word_document 生成并登记产物；把 Markdown/纯文本正文作为参数交给工具，由工具转换为 Word 原生标题、列表和表格。不要调用 bash、shell、python、execute，也不要把简单文档生成委派给 task 子任务。当前运行环境不提供命令执行后端。"""

RESOURCE_TOOL_READ_ONLY_SYSTEM_PROMPT = """上传文件属于任务资源，不是聊天上下文。
当用户询问上传文件时，优先使用资源工具。先检查或列出资源，再只读取所需的文本页或表格范围。不要假设上传文件内容已经直接出现在当前对话中。
读取 Word/docx 内容时，使用 inspect_resource、read_resource_text 或 read_resource_table。
当前任务没有提供文档生成工具；除非用户明确要求生成可下载文件，否则直接在最终回答中给出结论。不要调用 bash、shell、python、execute，也不要把简单资源读取或联网搜索委派给 task 子任务。"""

DEFAULT_TEXT_LIMIT = 40
MAX_TEXT_LIMIT = 200
DEFAULT_TABLE_LIMIT = 25
MAX_TABLE_LIMIT = 100
MAX_TABLE_COLUMNS = 80
MAX_WORD_MARKDOWN_CHARS = 120_000


@dataclass(frozen=True)
class ProvisionRequest:
    task_id: str
    workspace_root: Path


@dataclass(frozen=True)
class ExecutionRequest:
    tool_name: str
    input: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ExecutionError:
    code: str
    message: str
    retryable: bool = False

    def to_payload(self) -> dict[str, Any]:
        return {
            "code": self.code,
            "message": self.message,
            "retryable": self.retryable,
        }


@dataclass(frozen=True)
class ExecutionResult:
    ok: bool
    data: dict[str, Any] | list[Any] | str | int | float | bool | None = None
    error: ExecutionError | None = None

    @classmethod
    def success(cls, data: dict[str, Any] | list[Any] | str | int | float | bool | None):
        return cls(ok=True, data=data)

    @classmethod
    def failure(cls, code: str, message: str, *, retryable: bool = False):
        return cls(ok=False, error=ExecutionError(code, message, retryable))

    def to_payload(self) -> dict[str, Any]:
        if self.ok:
            return {"ok": True, "data": self.data}
        return {"ok": False, "error": self.error.to_payload() if self.error else None}

    def to_json(self) -> str:
        return json.dumps(self.to_payload(), ensure_ascii=False, default=str)


@dataclass(frozen=True)
class ResourceRecord:
    resource_id: str
    name: str
    path: Path
    format: str
    extension: str
    size_bytes: int
    digest: str

    def to_manifest_item(self) -> dict[str, Any]:
        return {
            "resource_id": self.resource_id,
            "name": self.name,
            "format": self.format,
            "extension": self.extension,
            "size_bytes": self.size_bytes,
            "digest": self.digest,
        }


@dataclass(frozen=True)
class ProvisionedResourceUnit:
    task_id: str
    workspace_root: Path
    task_workspace: Path
    upload_dir: Path
    resources: tuple[ResourceRecord, ...]
    run_id: str | None = None
    storage: Any | None = None


class LocalResourceExecutionAdapter:
    """In-process Provision/Execute adapter for uploaded task resources."""

    def __init__(
        self,
        *,
        task_id: str,
        workspace_root: Path,
        run_id: str | None = None,
        storage: Any | None = None,
    ):
        self.request = ProvisionRequest(task_id=task_id, workspace_root=workspace_root)
        self.run_id = run_id
        self.storage = storage

    def provision(self) -> ProvisionedResourceUnit:
        task_workspace = _resolve_task_workspace(self.request.workspace_root, self.request.task_id)
        upload_dir = task_workspace / "uploads"
        resources = tuple(_resource_records(self.request.task_id, upload_dir))
        return ProvisionedResourceUnit(
            task_id=self.request.task_id,
            workspace_root=self.request.workspace_root.resolve(),
            task_workspace=task_workspace,
            upload_dir=upload_dir,
            resources=resources,
            run_id=self.run_id,
            storage=self.storage,
        )

    def execute(self, request: ExecutionRequest) -> ExecutionResult:
        try:
            unit = self.provision()
        except ValueError as exc:
            return ExecutionResult.failure("invalid_task_workspace", str(exc))

        try:
            if request.tool_name == "list_uploaded_resources":
                return ExecutionResult.success(_list_uploaded_resources(unit))
            if request.tool_name == "inspect_resource":
                return _inspect_resource(unit, request.input)
            if request.tool_name == "read_resource_text":
                return _read_resource_text(unit, request.input)
            if request.tool_name == "read_resource_table":
                return _read_resource_table(unit, request.input)
            if request.tool_name == "create_word_document":
                return _create_word_document(unit, request.input)
            return ExecutionResult.failure(
                "unknown_tool",
                f"未知资源工具：{request.tool_name}",
            )
        except Exception as exc:  # pragma: no cover - safety net for tool-call isolation
            return ExecutionResult.failure("execution_error", str(exc), retryable=True)


def create_resource_tools(
    *,
    task_id: str,
    workspace_root: Path,
    run_id: str | None = None,
    storage: Any | None = None,
    include_artifact_tools: bool = True,
) -> list[BaseTool]:
    adapter = LocalResourceExecutionAdapter(
        task_id=task_id,
        workspace_root=workspace_root,
        run_id=run_id,
        storage=storage,
    )

    @tool
    def list_uploaded_resources() -> str:
        """列出当前任务下的上传资源，不读取文件正文。"""
        return adapter.execute(ExecutionRequest("list_uploaded_resources")).to_json()

    @tool
    def inspect_resource(resource: str) -> str:
        """按 resource_id 或文件名检查任务资源，并返回结构化元数据。"""
        return adapter.execute(
            ExecutionRequest("inspect_resource", {"resource": resource})
        ).to_json()

    @tool
    def read_resource_text(resource: str, page: int = 1, limit: int = DEFAULT_TEXT_LIMIT) -> str:
        """读取任务资源中的分页结构化文本块。"""
        return adapter.execute(
            ExecutionRequest(
                "read_resource_text",
                {"resource": resource, "page": page, "limit": limit},
            )
        ).to_json()

    @tool
    def read_resource_table(
        resource: str,
        sheet: str | None = None,
        range_ref: str | None = None,
        table_index: int | None = None,
        start_row: int = 1,
        limit: int = DEFAULT_TABLE_LIMIT,
    ) -> str:
        """从 Excel 或 Word 资源中读取表格或范围，并返回结构化行数据。"""
        return adapter.execute(
            ExecutionRequest(
                "read_resource_table",
                {
                    "resource": resource,
                    "sheet": sheet,
                    "range_ref": range_ref,
                    "table_index": table_index,
                    "start_row": start_row,
                    "limit": limit,
                },
            )
        ).to_json()

    tools: list[BaseTool] = [
        list_uploaded_resources,
        inspect_resource,
        read_resource_text,
        read_resource_table,
    ]

    if include_artifact_tools and run_id and storage is not None:
        @tool
        def create_word_document(filename: str, markdown: str) -> str:
            """从 Markdown/纯文本内容生成当前 run 的可下载 Word .docx 产物，表格会转换为 Word 原生表格。"""
            return adapter.execute(
                ExecutionRequest(
                    "create_word_document",
                    {"filename": filename, "markdown": markdown},
                )
            ).to_json()

        tools.append(create_word_document)

    return tools


def build_resource_manifest(task_id: str, workspace_root: Path) -> dict[str, Any]:
    unit = LocalResourceExecutionAdapter(task_id=task_id, workspace_root=workspace_root).provision()
    return _list_uploaded_resources(unit)


def format_resource_manifest_message(
    manifest: dict[str, Any],
    *,
    include_artifact_tools: bool = True,
) -> str:
    resources = manifest.get("resources") if isinstance(manifest, dict) else []
    if not resources:
        return ""
    return "\n".join(
        [
            (
                RESOURCE_TOOL_SYSTEM_PROMPT
                if include_artifact_tools
                else RESOURCE_TOOL_READ_ONLY_SYSTEM_PROMPT
            ),
            "当前任务资源清单：",
            json.dumps(
                {
                    "schema_version": 1,
                    "resources": resources,
                },
                ensure_ascii=False,
            ),
        ]
    )


def _resolve_task_workspace(workspace_root: Path, task_id: str) -> Path:
    root = workspace_root.resolve()
    task_workspace = (root / task_id).resolve()
    if task_workspace != root and root not in task_workspace.parents:
        raise ValueError("任务工作区超出资源根目录")
    return task_workspace


def _resource_records(task_id: str, upload_dir: Path) -> list[ResourceRecord]:
    if not upload_dir.exists():
        return []
    records: list[ResourceRecord] = []
    for path in sorted(upload_dir.iterdir()):
        if not path.is_file() or path.suffix.lower() not in UPLOAD_FORMATS:
            continue
        resource_ref = build_upload_resource_ref(
            session_id=task_id,
            filename=path.name,
            size_bytes=path.stat().st_size,
            digest=file_sha256(path),
            media_type=source_format_for_upload(path),
        )
        records.append(
            ResourceRecord(
                resource_id=resource_ref.id,
                name=path.name,
                path=path.resolve(),
                format=source_format_for_upload(path),
                extension=path.suffix.lower(),
                size_bytes=path.stat().st_size,
                digest=resource_ref.digest or "",
            )
        )
    return records


def _list_uploaded_resources(unit: ProvisionedResourceUnit) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "task_id": unit.task_id,
        "resources": [resource.to_manifest_item() for resource in unit.resources],
    }


def _resolve_resource(unit: ProvisionedResourceUnit, raw_resource: Any) -> ResourceRecord | None:
    resource = str(raw_resource or "").strip()
    if not resource:
        return None
    by_id = {item.resource_id: item for item in unit.resources}
    if resource in by_id:
        return by_id[resource]
    if resource.startswith(f"upload:{unit.task_id}:"):
        resource = resource.rsplit(":", 1)[-1]
    if Path(resource).name != resource:
        return None
    resource_key = resource.casefold()
    for item in unit.resources:
        if item.name.casefold() == resource_key:
            return item
    return None


def _inspect_resource(unit: ProvisionedResourceUnit, input_data: dict[str, Any]) -> ExecutionResult:
    resource = _resolve_resource(unit, input_data.get("resource"))
    if resource is None:
        return ExecutionResult.failure("resource_not_found", "未找到当前任务中的上传资源")
    data = {
        **resource.to_manifest_item(),
        "structure": _inspect_structure(resource),
    }
    return ExecutionResult.success(data)


def _inspect_structure(resource: ResourceRecord) -> dict[str, Any]:
    if resource.format in {"markdown", "text"}:
        text = _read_text_file(resource.path)
        lines = text.splitlines()
        return {
            "kind": resource.format,
            "line_count": len(lines),
            "char_count": len(text),
            "preview_lines": lines[: min(5, len(lines))],
        }
    if resource.format == "json":
        value = _load_json_file(resource.path)
        return {
            "kind": "json",
            "json_type": type(value).__name__,
            "top_level_keys": list(value.keys())[:50] if isinstance(value, dict) else None,
            "length": len(value) if isinstance(value, (dict, list)) else None,
        }
    if resource.format == "word":
        return _inspect_docx(resource.path)
    if resource.format == "excel":
        return _inspect_excel(resource.path)
    return {"kind": resource.format}


def _read_resource_text(unit: ProvisionedResourceUnit, input_data: dict[str, Any]) -> ExecutionResult:
    resource = _resolve_resource(unit, input_data.get("resource"))
    if resource is None:
        return ExecutionResult.failure("resource_not_found", "未找到当前任务中的上传资源")
    page = _positive_int(input_data.get("page"), 1)
    limit = min(_positive_int(input_data.get("limit"), DEFAULT_TEXT_LIMIT), MAX_TEXT_LIMIT)

    if resource.format in {"markdown", "text"}:
        return ExecutionResult.success(_paginate_lines(_read_text_file(resource.path), page, limit, resource))
    if resource.format == "json":
        text = json.dumps(_load_json_file(resource.path), ensure_ascii=False, indent=2)
        return ExecutionResult.success(_paginate_lines(text, page, limit, resource))
    if resource.format == "word":
        blocks = _docx_blocks(resource.path)
        return ExecutionResult.success(_paginate_blocks(blocks, page, limit, resource))
    if resource.format == "excel":
        return ExecutionResult.failure(
            "unsupported_operation",
            "Excel 资源请使用 read_resource_table 按 sheet 或 range 读取",
        )
    return ExecutionResult.failure("unsupported_resource_format", f"暂不支持读取 {resource.format}")


def _read_resource_table(unit: ProvisionedResourceUnit, input_data: dict[str, Any]) -> ExecutionResult:
    resource = _resolve_resource(unit, input_data.get("resource"))
    if resource is None:
        return ExecutionResult.failure("resource_not_found", "未找到当前任务中的上传资源")
    limit = min(_positive_int(input_data.get("limit"), DEFAULT_TABLE_LIMIT), MAX_TABLE_LIMIT)

    if resource.format == "excel":
        return _read_excel_table(
            resource,
            sheet=input_data.get("sheet"),
            range_ref=input_data.get("range_ref"),
            start_row=_positive_int(input_data.get("start_row"), 1),
            limit=limit,
        )
    if resource.format == "word":
        table_index = _positive_int(input_data.get("table_index"), 0, allow_zero=True)
        return ExecutionResult.success(_read_docx_table(resource, table_index=table_index, limit=limit))
    return ExecutionResult.failure(
        "unsupported_operation",
        "该资源格式不包含可读取的表格范围",
    )


def _create_word_document(
    unit: ProvisionedResourceUnit,
    input_data: dict[str, Any],
) -> ExecutionResult:
    if unit.run_id is None or unit.storage is None:
        return ExecutionResult.failure(
            "artifact_context_unavailable",
            "当前运行缺少产物登记上下文，无法生成可下载 Word 文件",
        )

    raw_markdown = input_data.get("markdown")
    markdown = raw_markdown if isinstance(raw_markdown, str) else str(raw_markdown or "")
    markdown = markdown.strip()
    if not markdown:
        return ExecutionResult.failure(
            "empty_document",
            "create_word_document 需要提供 markdown 正文",
        )
    if len(markdown) > MAX_WORD_MARKDOWN_CHARS:
        return ExecutionResult.failure(
            "document_too_large",
            f"Word 正文过长，请控制在 {MAX_WORD_MARKDOWN_CHARS} 字以内",
        )

    artifact_name = _word_artifact_name(input_data.get("filename"))
    artifact_dir = unit.storage.run_artifact_dir(unit.task_id, unit.run_id)
    artifact_dir.mkdir(parents=True, exist_ok=True)
    artifact_path = (artifact_dir / artifact_name).resolve()
    if artifact_dir.resolve() not in artifact_path.parents:
        return ExecutionResult.failure("invalid_artifact_path", "Word 产物路径无效")

    _write_docx_from_markdown(artifact_path, markdown)
    unit.storage.record_run_artifact(unit.task_id, unit.run_id, artifact_name)
    relative_path = artifact_path.relative_to(unit.task_workspace)
    return ExecutionResult.success(
        {
            "schema_version": 1,
            "artifact_name": artifact_name,
            "artifact_type": "word",
            "run_id": unit.run_id,
            "relative_path": relative_path.as_posix(),
            "download_url": (
                f"/api/tasks/{unit.task_id}/runs/{unit.run_id}/artifacts/{artifact_name}"
            ),
            "size_bytes": artifact_path.stat().st_size,
        }
    )


def _word_artifact_name(raw_filename: Any) -> str:
    candidate = str(raw_filename or "generated-document.docx").strip()
    if not candidate:
        candidate = "generated-document.docx"
    if Path(candidate).suffix.lower() != ".docx":
        candidate = f"{Path(candidate).stem or candidate}.docx"
    return normalize_artifact_name(Path(candidate).name)


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _load_json_file(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _paginate_lines(text: str, page: int, limit: int, resource: ResourceRecord) -> dict[str, Any]:
    lines = text.splitlines()
    start = (page - 1) * limit
    end = start + limit
    return {
        "schema_version": 1,
        "resource": resource.to_manifest_item(),
        "page": page,
        "limit": limit,
        "total_items": len(lines),
        "total_pages": max(1, math.ceil(len(lines) / limit)) if limit else 1,
        "items": [
            {"line": index + 1, "text": line}
            for index, line in enumerate(lines[start:end], start=start)
        ],
    }


def _paginate_blocks(
    blocks: list[dict[str, Any]], page: int, limit: int, resource: ResourceRecord
) -> dict[str, Any]:
    start = (page - 1) * limit
    end = start + limit
    return {
        "schema_version": 1,
        "resource": resource.to_manifest_item(),
        "page": page,
        "limit": limit,
        "total_items": len(blocks),
        "total_pages": max(1, math.ceil(len(blocks) / limit)) if limit else 1,
        "items": blocks[start:end],
    }


def _write_docx_from_markdown(path: Path, markdown: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    lines = markdown.splitlines()
    title_written = False
    wrote_content = False
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        if _is_markdown_fence(line):
            index += 1
            continue
        if _is_markdown_horizontal_rule(line):
            index += 1
            continue
        table_rows, consumed = _consume_markdown_table(lines, index)
        if table_rows is not None:
            _add_docx_table(document, table_rows)
            wrote_content = True
            index += consumed
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = _plain_markdown_text(heading.group(2))
            paragraph = document.add_heading(text, level=level)
            if not title_written and level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_written = True
            wrote_content = True
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            document.add_paragraph(_plain_markdown_text(bullet.group(1)), style="List Bullet")
            wrote_content = True
            index += 1
            continue
        numbered = re.match(r"^\d+[.)、]\s+(.+)$", line)
        if numbered:
            document.add_paragraph(_plain_markdown_text(numbered.group(1)), style="List Number")
            wrote_content = True
            index += 1
            continue
        document.add_paragraph(_plain_markdown_text(line))
        wrote_content = True
        index += 1

    if not wrote_content:
        document.add_paragraph(markdown)
    document.save(str(path))


def _consume_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]] | None, int]:
    if start + 1 >= len(lines):
        return None, 0
    header = _split_markdown_table_row(lines[start])
    separator = _split_markdown_table_row(lines[start + 1])
    if not header or not _is_markdown_table_separator(separator):
        return None, 0

    rows = [header]
    index = start + 2
    while index < len(lines):
        row = _split_markdown_table_row(lines[index])
        if not row:
            break
        rows.append(row)
        index += 1
    return _normalize_table_rows(rows), index - start


def _split_markdown_table_row(line: str) -> list[str] | None:
    stripped = line.strip()
    if "|" not in stripped:
        return None
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells = [cell.strip() for cell in stripped.split("|")]
    return cells if len(cells) >= 2 else None


def _is_markdown_table_separator(cells: list[str] | None) -> bool:
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _normalize_table_rows(rows: list[list[str]]) -> list[list[str]]:
    column_count = max((len(row) for row in rows), default=0)
    return [row + [""] * (column_count - len(row)) for row in rows]


def _add_docx_table(document: Any, rows: list[list[str]]) -> None:
    if not rows or not rows[0]:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, cell_text in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            run = paragraph.add_run(_plain_markdown_text(cell_text))
            if row_index == 0:
                run.bold = True


def _is_markdown_horizontal_rule(line: str) -> bool:
    return bool(re.fullmatch(r"[-*_]{3,}", line))


def _is_markdown_fence(line: str) -> bool:
    return line.startswith("```") or line.startswith("~~~")


def _plain_markdown_text(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text.strip()


def _inspect_docx(path: Path) -> dict[str, Any]:
    blocks = _docx_blocks(path)
    table_count = sum(1 for block in blocks if block["type"] == "table")
    headings = [
        {
            "order": block["order"],
            "level": block.get("level"),
            "text": block.get("text"),
        }
        for block in blocks
        if block["type"] == "heading"
    ][:50]
    return {
        "kind": "word",
        "block_count": len(blocks),
        "table_count": table_count,
        "headings": headings,
    }


def _docx_blocks(path: Path) -> list[dict[str, Any]]:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    blocks: list[dict[str, Any]] = []
    table_index = 0
    paragraph_index = 0

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            block_type = "paragraph"
            level: int | None = None
            if style_name.lower().startswith("heading"):
                block_type = "heading"
                level = _heading_level(style_name)
            elif "list" in style_name.lower():
                block_type = "list_item"
            item: dict[str, Any] = {
                "order": len(blocks),
                "type": block_type,
                "paragraph_index": paragraph_index,
                "style": style_name,
                "text": text,
            }
            if level is not None:
                item["level"] = level
            blocks.append(item)
            paragraph_index += 1
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            blocks.append(
                {
                    "order": len(blocks),
                    "type": "table",
                    "table_index": table_index,
                    "row_count": len(rows),
                    "column_count": max((len(row) for row in rows), default=0),
                    "rows": rows[: min(len(rows), DEFAULT_TABLE_LIMIT)],
                    "truncated": len(rows) > DEFAULT_TABLE_LIMIT,
                }
            )
            table_index += 1
    return blocks


def _read_docx_table(resource: ResourceRecord, *, table_index: int, limit: int) -> dict[str, Any]:
    tables = [block for block in _docx_blocks(resource.path) if block["type"] == "table"]
    if table_index >= len(tables):
        return {
            "schema_version": 1,
            "resource": resource.to_manifest_item(),
            "table_index": table_index,
            "rows": [],
            "error": {"code": "table_not_found", "message": "Word 表格不存在"},
        }
    table = tables[table_index]
    rows = table.get("rows", [])
    return {
        "schema_version": 1,
        "resource": resource.to_manifest_item(),
        "table_index": table_index,
        "row_count": table.get("row_count", len(rows)),
        "column_count": table.get("column_count", 0),
        "rows": rows[:limit],
        "truncated": len(rows) > limit or bool(table.get("truncated")),
    }


def _inspect_excel(path: Path) -> dict[str, Any]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=False)
    sheets = []
    for sheet in workbook.worksheets:
        sheets.append(
            {
                "name": sheet.title,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "merged_ranges": [str(item) for item in list(sheet.merged_cells.ranges)[:50]],
                "header_guess": _guess_header(sheet),
            }
        )
    workbook.close()
    return {
        "kind": "excel",
        "sheet_count": len(sheets),
        "sheets": sheets,
    }


def _read_excel_table(
    resource: ResourceRecord,
    *,
    sheet: Any,
    range_ref: Any,
    start_row: int,
    limit: int,
) -> ExecutionResult:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter, range_boundaries

    workbook = load_workbook(resource.path, data_only=True, read_only=False)
    try:
        sheet_name = str(sheet).strip() if sheet else workbook.sheetnames[0]
        if sheet_name not in workbook.sheetnames:
            return ExecutionResult.success(
                {
                    "schema_version": 1,
                    "resource": resource.to_manifest_item(),
                    "sheet": sheet_name,
                    "rows": [],
                    "error": {"code": "sheet_not_found", "message": "Excel sheet 不存在"},
                }
            )
        worksheet = workbook[sheet_name]
        if range_ref:
            try:
                min_col, min_row, max_col, max_row = range_boundaries(str(range_ref))
            except ValueError:
                return ExecutionResult.failure(
                    "invalid_range",
                    "Excel range_ref 无效，请使用类似 A1:D20 的单元格范围",
                )
            if (
                min_col is None
                or min_row is None
                or max_col is None
                or max_row is None
                or min_col < 1
                or min_row < 1
                or max_col < min_col
                or max_row < min_row
            ):
                return ExecutionResult.failure(
                    "invalid_range",
                    "Excel range_ref 无效，请使用类似 A1:D20 的单元格范围",
                )
            max_row = min(max_row, min_row + limit - 1)
        else:
            min_row = start_row
            max_row = min(worksheet.max_row, start_row + limit - 1)
            min_col = 1
            max_col = min(worksheet.max_column, MAX_TABLE_COLUMNS)
        worksheet_max_row = worksheet.max_row
        rows = []
        for row in worksheet.iter_rows(
            min_row=min_row,
            max_row=max_row,
            min_col=min_col,
            max_col=max_col,
            values_only=True,
        ):
            rows.append([_cell_value(value) for value in row])
        return ExecutionResult.success(
            {
                "schema_version": 1,
                "resource": resource.to_manifest_item(),
                "sheet": sheet_name,
                "range": f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{max_row}",
                "start_row": min_row,
                "row_count": len(rows),
                "column_count": max_col - min_col + 1,
                "rows": rows,
                "truncated": max_row < worksheet_max_row if not range_ref else False,
            }
        )
    finally:
        workbook.close()


def _guess_header(sheet: Any) -> dict[str, Any] | None:
    for row_index, row in enumerate(
        sheet.iter_rows(min_row=1, max_row=min(sheet.max_row, 20), values_only=True),
        start=1,
    ):
        values = [_cell_value(value) for value in row]
        non_empty = [value for value in values if value not in {"", None}]
        if len(non_empty) >= 2:
            return {"row": row_index, "values": values[:MAX_TABLE_COLUMNS]}
    return None


def _cell_value(value: Any) -> Any:
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return value


def _heading_level(style_name: str) -> int | None:
    parts = style_name.rsplit(" ", 1)
    if len(parts) != 2:
        return None
    try:
        return int(parts[1])
    except ValueError:
        return None


def _positive_int(value: Any, fallback: int, *, allow_zero: bool = False) -> int:
    try:
        number = int(value)
    except (TypeError, ValueError):
        return fallback
    minimum = 0 if allow_zero else 1
    return number if number >= minimum else fallback
