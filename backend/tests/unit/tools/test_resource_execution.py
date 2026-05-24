from __future__ import annotations

import json

from docx import Document
from openpyxl import Workbook

from app.execution.resources import (
    ExecutionRequest,
    LocalResourceExecutionAdapter,
    build_resource_manifest,
    create_resource_tools,
    format_resource_manifest_message,
)
from tests.fakes import InMemoryTaskStorage


def _task_upload_dir(tmp_path, task_id: str):
    upload_dir = tmp_path / task_id / "uploads"
    upload_dir.mkdir(parents=True)
    return upload_dir


def _write_docx(path):
    document = Document()
    document.add_heading("项目概览", level=1)
    document.add_paragraph("这是一段 Word 正文。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "预算"
    table.cell(1, 1).text = "100"
    document.save(path)


def _write_xlsx(path):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "明细"
    sheet.append(["名称", "金额"])
    sheet.append(["A", 10])
    sheet.append(["B", 20])
    workbook.save(path)


class TestLocalResourceExecutionAdapter:
    def test_lists_manifest_without_reading_contents(self, tmp_path):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        (upload_dir / "notes.txt").write_text("secret body", encoding="utf-8")

        manifest = build_resource_manifest(task_id, tmp_path)

        assert manifest["resources"][0]["name"] == "notes.txt"
        assert manifest["resources"][0]["format"] == "text"
        assert "secret body" not in json.dumps(manifest, ensure_ascii=False)

    def test_reads_text_resource_as_paginated_json(self, tmp_path):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        (upload_dir / "notes.txt").write_text("alpha\nbeta\ngamma", encoding="utf-8")
        adapter = LocalResourceExecutionAdapter(task_id=task_id, workspace_root=tmp_path)

        result = adapter.execute(
            ExecutionRequest(
                "read_resource_text",
                {"resource": "notes.txt", "page": 1, "limit": 2},
            )
        ).to_payload()

        assert result["ok"] is True
        assert result["data"]["total_items"] == 3
        assert [item["text"] for item in result["data"]["items"]] == ["alpha", "beta"]

    def test_inspects_docx_structure_and_reads_table(self, tmp_path):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        _write_docx(upload_dir / "brief.docx")
        adapter = LocalResourceExecutionAdapter(task_id=task_id, workspace_root=tmp_path)

        inspected = adapter.execute(
            ExecutionRequest("inspect_resource", {"resource": "brief.docx"})
        ).to_payload()
        table = adapter.execute(
            ExecutionRequest("read_resource_table", {"resource": "brief.docx", "table_index": 0})
        ).to_payload()

        assert inspected["ok"] is True
        assert inspected["data"]["structure"]["kind"] == "word"
        assert inspected["data"]["structure"]["headings"][0]["text"] == "项目概览"
        assert table["ok"] is True
        assert table["data"]["rows"][1] == ["预算", "100"]

    def test_inspects_xlsx_structure_and_reads_sheet_range(self, tmp_path):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        _write_xlsx(upload_dir / "data.xlsx")
        adapter = LocalResourceExecutionAdapter(task_id=task_id, workspace_root=tmp_path)

        inspected = adapter.execute(
            ExecutionRequest("inspect_resource", {"resource": "data.xlsx"})
        ).to_payload()
        table = adapter.execute(
            ExecutionRequest(
                "read_resource_table",
                {"resource": "data.xlsx", "sheet": "明细", "range_ref": "A1:B2"},
            )
        ).to_payload()

        assert inspected["ok"] is True
        assert inspected["data"]["structure"]["sheets"][0]["header_guess"]["values"] == ["名称", "金额"]
        assert table["ok"] is True
        assert table["data"]["rows"] == [["名称", "金额"], ["A", 10]]

    def test_read_xlsx_invalid_range_is_non_retryable_and_closes_workbook(
        self, tmp_path, monkeypatch
    ):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        (upload_dir / "data.xlsx").write_bytes(b"placeholder")
        closed = {"value": False}

        class FakeWorksheet:
            max_row = 2
            max_column = 2

            def iter_rows(self, **kwargs):
                raise AssertionError("invalid range should stop before reading rows")

        class FakeWorkbook:
            sheetnames = ["明细"]

            def __getitem__(self, sheet_name):
                assert sheet_name == "明细"
                return FakeWorksheet()

            def close(self):
                closed["value"] = True

        def fake_range_boundaries(range_ref):
            assert range_ref == "not-a-range"
            raise ValueError("invalid range")

        monkeypatch.setattr("openpyxl.load_workbook", lambda *args, **kwargs: FakeWorkbook())
        monkeypatch.setattr("openpyxl.utils.range_boundaries", fake_range_boundaries)
        adapter = LocalResourceExecutionAdapter(task_id=task_id, workspace_root=tmp_path)

        result = adapter.execute(
            ExecutionRequest(
                "read_resource_table",
                {"resource": "data.xlsx", "sheet": "明细", "range_ref": "not-a-range"},
            )
        ).to_payload()

        assert result == {
            "ok": False,
            "error": {
                "code": "invalid_range",
                "message": "Excel range_ref 无效，请使用类似 A1:D20 的单元格范围",
                "retryable": False,
            },
        }
        assert closed["value"] is True

    def test_read_xlsx_incomplete_range_is_non_retryable(self, tmp_path):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        _write_xlsx(upload_dir / "data.xlsx")
        adapter = LocalResourceExecutionAdapter(task_id=task_id, workspace_root=tmp_path)

        result = adapter.execute(
            ExecutionRequest(
                "read_resource_table",
                {"resource": "data.xlsx", "sheet": "明细", "range_ref": "A:B"},
            )
        ).to_payload()

        assert result["ok"] is False
        assert result["error"]["code"] == "invalid_range"
        assert result["error"]["retryable"] is False

    def test_missing_resource_returns_tool_call_error_payload(self, tmp_path):
        adapter = LocalResourceExecutionAdapter(task_id="task-1", workspace_root=tmp_path)

        result = adapter.execute(
            ExecutionRequest("inspect_resource", {"resource": "../outside.txt"})
        ).to_payload()

        assert result == {
            "ok": False,
            "error": {
                "code": "resource_not_found",
                "message": "未找到当前任务中的上传资源",
                "retryable": False,
            },
        }

    def test_langchain_tools_delegate_to_execution_adapter(self, tmp_path):
        task_id = "task-1"
        upload_dir = _task_upload_dir(tmp_path, task_id)
        (upload_dir / "data.json").write_text('{"name": "MyAgent"}', encoding="utf-8")
        tools = {tool.name: tool for tool in create_resource_tools(task_id=task_id, workspace_root=tmp_path)}

        result = json.loads(tools["inspect_resource"].invoke({"resource": "data.json"}))

        assert result["ok"] is True
        assert result["data"]["format"] == "json"
        assert result["data"]["structure"]["top_level_keys"] == ["name"]

    def test_create_word_document_writes_run_artifact(self, tmp_path):
        storage = InMemoryTaskStorage(tmp_path)
        state = storage.create_task(message=None, model="deepseek-v4-flash")
        run_result = storage.start_run(
            state.task_id,
            message="生成 Word 总结",
            model="deepseek-v4-flash",
            expected_statuses={"idle"},
        )
        assert run_result is not None
        _, run_id = run_result
        tools = {
            tool.name: tool
            for tool in create_resource_tools(
                task_id=state.task_id,
                workspace_root=tmp_path,
                run_id=run_id,
                storage=storage,
            )
        }

        result = json.loads(
            tools["create_word_document"].invoke(
                {
                    "filename": "技术参数总结.docx",
                    "markdown": (
                        "```markdown\n"
                        "# 技术参数总结\n\n"
                        "| 项目 | 内容 |\n"
                        "| --- | --- |\n"
                        "| 屏幕规格 | ≥30英寸 |\n"
                        "| 预算 | 140万元 |\n\n"
                        "- 符合 DICOM 3.14 标准\n"
                        "```"
                    ),
                }
            )
        )

        assert result["ok"] is True
        assert result["data"]["artifact_name"] == "技术参数总结.docx"
        artifact_path = storage.resolve_run_artifact(
            state.task_id,
            run_id,
            "技术参数总结.docx",
        )
        assert artifact_path.exists()
        document = Document(str(artifact_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "技术参数总结" in text
        assert len(document.tables) == 1
        assert document.tables[0].cell(0, 0).text == "项目"
        assert document.tables[0].cell(1, 0).text == "屏幕规格"
        assert document.tables[0].cell(1, 1).text == "≥30英寸"
        assert all("| --- |" not in paragraph.text for paragraph in document.paragraphs)
        assert all("```" not in paragraph.text for paragraph in document.paragraphs)
        assert "符合 DICOM 3.14 标准" in text

    def test_resource_manifest_message_omits_empty_manifest(self, tmp_path):
        assert format_resource_manifest_message(build_resource_manifest("task-1", tmp_path)) == ""
